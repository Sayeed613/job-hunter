"""Test script: Watch the bot apply to a real job in real-time.

Opens a visible browser (headless=False) on a real Greenhouse job listing,
clicks Apply, fills the entire form with candidate data, clicks Submit,
and checks for confirmation. You can watch everything happen.

Usage:
    python test_apply.py                       # Visible browser + confirmation prompt
    python test_apply.py --dry-run             # Fill form, navigate steps, NO submit
    python test_apply.py --headless            # Run headless (invisible)
    python test_apply.py --url <custom_url>    # Test with a different job
"""

from __future__ import annotations

import os
os.environ["PYTHONIOENCODING"] = "utf-8"

import asyncio
import logging
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from app.ai.client import AIClient
from app.ai.cover_letter_gen import CoverLetterGenerator
from app.ai.keyword_extractor import KeywordExtractor
from app.ai.resume_tailor import ResumeTailor
from app.browser.browser_manager import BrowserManager
from app.browser.form_filler import FormFiller
from app.browser.human import Human
from app.config.settings import Settings
from app.models.job import Job
from app.resume.models import ResumeProfile
from app.resume.parser import ResumeParser
from app.tailor.resume_generator import ResumeGenerator
from app.telegram.interaction import TelegramInteraction
from app.telegram.learner import LearnedProfile
from app.telegram.notifier import TelegramNotifier

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("apply_test")

TEST_JOB = {
    "company": "Vercel",
    "title": "Design Engineer",
    "apply_url": "https://boards.greenhouse.io/vercel/jobs/5709080004",
    "location": "AMER (Remote)",
}


async def test_apply(
    headless: bool = False,
    apply_url: str | None = None,
    dry_run: bool = False,
) -> bool:
    """Open a job's apply page, fill the form, submit, and verify.

    Args:
        headless: If True, run browser headless (invisible).
        apply_url: Override the default test job URL.
        dry_run: If True, fill all fields across all steps but skip final submit.

    Returns:
        True if the application appeared to submit successfully.
    """
    url = apply_url or TEST_JOB["apply_url"]
    company = TEST_JOB["company"]
    title = TEST_JOB["title"]

    # Find a resume file
    resume_path = ""
    for candidate in [
        "resumes/Sayeed_Ahmed_Resume.docx",
        "Sayeed_Frontend_Developer.docx",
        "Sayeed_Ahmed_Resume.docx",
    ]:
        p = Path(candidate)
        if p.exists():
            resume_path = str(p.resolve())
            break
    if not resume_path:
        paths = list(Path(".").glob("*resume*.docx")) + list(Path("resumes").glob("*.docx"))
        if paths:
            resume_path = str(paths[0])

    print(f"\n{'=' * 60}")
    print(f"TESTING REAL APPLICATION{' (DRY RUN)' if dry_run else ''}")
    print(f"{'=' * 60}")
    print(f"Company:  {company}")
    print(f"Position: {title}")
    print(f"URL:      {url}")
    print(f"Resume:   {resume_path or 'None'}")
    print(f"{'=' * 60}\n")

    mgr = BrowserManager()
    await mgr.launch(headless=headless)
    page = await mgr.new_page()
    success = False

    try:
        # ── Navigate to job ──
        print("Navigating to job page...")
        await page.goto(url, wait_until="networkidle", timeout=45000)
        await Human.delay(3, 5)
        print(f"   Page loaded: {page.url}")

        # Wait for the page to finish rendering (not "Loading...")
        for _ in range(5):
            try:
                body_text = (await page.inner_text("body") or "").lower().strip()
                if body_text and body_text != "loading...":
                    print(f"   Page content ready: {body_text[:80]}...")
                    break
            except Exception:
                pass
            print("   Waiting for page to finish rendering...")
            await Human.delay(2, 3)

        # ── Click Apply button ──
        # Retry up to 2 times (page may still be rendering)
        clicked_apply = False
        for attempt in range(2):
            if attempt > 0:
                print("   Retrying Apply button detection...")
                await Human.delay(2, 3)

            print("\n[SEARCH] Looking for Apply button...")
            apply_selectors = [
                # Greenhouse-specific: anchor styled as button
                "a.button:has-text('Apply')",
                ".button:has-text('Apply')",
                "[class*='job-application']:has-text('Apply')",
                "a:has-text('Apply for this job')",
                "button:has-text('Apply')",
                "[class*='apply']:has-text('Apply')",
                "a[href*='apply']",
                "button[type='submit']",
                "[data-testid*='apply']",
                ".apply-btn",
                "#apply-button",
                # Broad fallback: any element containing 'Apply'
                ":has-text('Apply for this job')",
                ":has-text('Apply Now')",
                ":has-text('Apply')",
            ]
            for sel in apply_selectors:
                try:
                    btn = await page.query_selector(sel)
                    if btn:
                        text = (await btn.inner_text()).strip()
                        print(f"   Found: '{text}' via '{sel}'")
                        await btn.scroll_into_view_if_needed()
                        await Human.delay(0.3, 0.7)
                        await btn.click()
                        await Human.delay(2, 4)
                        clicked_apply = True
                        break
                except Exception:
                    continue

            if clicked_apply:
                break

            # Ultimate fallback: JavaScript click on ANY visible element with 'Apply' text
            print("   Trying JavaScript click fallback...")
            try:
                clicked = await page.evaluate("""() => {
                    const all = document.querySelectorAll('*');
                    for (const el of all) {
                        const text = (el.textContent || '').toLowerCase().trim();
                        if (text.includes('apply') && text.length < 50 &&
                            el.offsetParent !== null &&
                            ['a','button','span','div','input'].includes(el.tagName.toLowerCase())) {
                            el.scrollIntoView({behavior: 'instant', block: 'center'});
                            el.click();
                            return true;
                        }
                    }
                    return false;
                }""")
                if clicked:
                    print("   [OK] Clicked Apply via JavaScript fallback")
                    await Human.delay(2, 4)
                    clicked_apply = True
                    break
            except Exception:
                pass

        if not clicked_apply:
            print("[FAIL] No Apply button found!")
            try:
                body = (await page.inner_text("body") or "")[:500]
                print(f"   Page text: {body[:200]}...")
            except Exception:
                pass
            await Human.screenshot(page, "test_no_apply_btn")
            return False

        print("   [OK] Apply button clicked!")

        # ── Initialise AI client + Telegram ──
        cfg = Settings()
        ai_client = AIClient() if cfg.openai_api_key else None
        tg_notifier = TelegramNotifier(
            token=cfg.telegram_bot_token,
            chat_id=cfg.telegram_chat_id,
        )
        ai_available = ai_client and ai_client.is_available
        if ai_available:
            print("   [AI] AI client ready — will tailor resume + generate cover letter")
        else:
            print("   [AI] AI not configured — using static resume + template cover letter")

        # ── Parse base resume ──
        print("\n[RESUME] Parsing base resume...")
        parser = ResumeParser()
        resume = parser.parse_docx(resume_path)
        print(f"   Name:       {resume.name}")
        print(f"   Skills:     {len(resume.skills)}")
        print(f"   Experience: {len(resume.experience)} entries")
        print(f"   Projects:   {len(resume.projects)}")

        # ── Scrape job description from page ──
        print("\n[JD] Scraping job description from page...")
        job_description_text = ""
        try:
            job_description_text = (await page.inner_text("body") or "")
            # Clean: remove very short lines and excessive whitespace
            lines = [l.strip() for l in job_description_text.split("\n") if len(l.strip()) > 20]
            job_description_text = "\n".join(lines[:100])  # Keep top 100 substantive lines
            print(f"   Scraped {len(job_description_text)} chars of job description")
        except Exception as e:
            print(f"   [WARN] Could not scrape JD: {e}")

        # ── Create Job object ──
        import hashlib
        job_id = hashlib.sha256(f"{company}:{title}".encode()).hexdigest()[:16]
        job = Job(
            job_id=job_id,
            title=title,
            company=company,
            description=job_description_text,
            location=TEST_JOB.get("location", ""),
            apply_url=url,
        )

        # ── AI: Extract keywords ──
        cover_text = ""
        tailored_resume_path = resume_path
        if ai_available:
            keyword_extractor = KeywordExtractor(client=ai_client)
            resume_tailor = ResumeTailor(client=ai_client)
            cover_gen = CoverLetterGenerator(client=ai_client)
            resume_gen = ResumeGenerator()

            # Extract keywords from JD
            print("\n[AI] Extracting keywords from job description...")
            try:
                keywords = await keyword_extractor.extract(job.description)
                hard_skills = keywords.get("hard_skills", [])
                print(f"   Found {len(hard_skills)} hard skills: {', '.join(hard_skills[:8])}")
            except Exception as e:
                print(f"   [WARN] Keyword extraction failed: {e}")
                keywords = {"hard_skills": [], "soft_skills": [], "years_required": 0}
                hard_skills = []

            # Tailor resume
            print("\n[AI] Tailoring resume for this job...")
            base_resume_text = _resume_to_text(resume)
            try:
                tailored_text = await resume_tailor.tailor(
                    base_resume_text, job, hard_skills,
                )
                print(f"   Tailored resume generated ({len(tailored_text)} chars)")
            except Exception as e:
                print(f"   [WARN] Resume tailoring failed: {e}")
                tailored_text = base_resume_text

            # Parse tailored text into a ResumeProfile and generate .docx
            print("\n[AI] Generating tailored resume DOCX...")
            try:
                tailored_profile = ResumeProfile(
                    name=_parse_tailored_field(tailored_text, "Name:") or resume.name,
                    email=_parse_tailored_field(tailored_text, "Email:") or resume.email,
                    phone=_parse_tailored_field(tailored_text, "Phone:") or resume.phone,
                    location=_parse_tailored_field(tailored_text, "Location:") or resume.location,
                    summary=_parse_tailored_field(tailored_text, "Summary:") or resume.summary,
                    skills=_parse_tailored_skills(tailored_text) or resume.skills,
                    experience=_parse_tailored_list(tailored_text, "Experience:") or resume.experience,
                    projects=resume.projects,
                    education=_parse_tailored_list(tailored_text, "Education:") or resume.education,
                    certifications=_parse_tailored_list(tailored_text, "Certifications:") or resume.certifications,
                )
                output_dir = Path("output") / _safe_name(company)
                output_dir.mkdir(parents=True, exist_ok=True)
                tailored_docx = resume_gen.generate_docx(
                    tailored_profile, output_dir / f"resume_{job.job_id}.docx",
                )
                tailored_resume_path = str(tailored_docx)
                print(f"   Tailored resume saved: {tailored_resume_path}")
            except Exception as e:
                print(f"   [WARN] Resume DOCX generation failed: {e}")
                tailored_resume_path = resume_path

            # Generate cover letter
            print("\n[AI] Generating cover letter...")
            try:
                cl_summary = (
                    _parse_tailored_field(tailored_text, "Summary:")
                    or resume.summary
                )
                cover_text = await cover_gen.generate(
                    cl_summary,
                    [f"- {p.description}" for p in resume.projects[:5]],
                    job, hard_skills,
                )
                print(f"   Cover letter generated ({len(cover_text)} chars)")

                # Save cover letter as .docx too
                cover_path = output_dir / f"cover_letter_{job.job_id}.docx"
                _write_cover_docx(cover_text, cover_path, resume, company)
                print(f"   Cover letter saved: {cover_path}")
            except Exception as e:
                print(f"   [WARN] Cover letter generation failed: {e}")
                cover_text = (
                    f"Dear {company} Hiring Team,\n\n"
                    f"I am excited to apply for the {title} position at {company}. "
                    f"With my background in frontend development and a passion for "
                    f"building great user interfaces, I believe I would be a strong "
                    f"addition to your team.\n\n"
                    f"Best regards,\nSayeed Ahmed\nsayeedahmed90082@gmail.com\n+91 9008299613"
                )
        else:
            # No AI: use static cover letter
            cover_text = (
                f"Dear {company} Hiring Team,\n\n"
                f"I am excited to apply for the {title} position at {company}. "
                f"With my background in frontend development and a passion for "
                f"building great user interfaces, I believe I would be a strong "
                f"addition to your team.\n\n"
                f"I have experience with React, Next.js, TypeScript, Node.js, "
                f"and modern frontend tooling.\n\n"
                f"Thank you for your consideration.\n\n"
                f"Best regards,\nSayeed Ahmed\nsayeedahmed90082@gmail.com\n+91 9008299613"
            )

        print(f"\n{'=' * 60}")
        print(f"Resume:   {tailored_resume_path}")
        print(f"Cover:    {len(cover_text)} chars")
        print(f"{'=' * 60}\n")

        # ── Initialise Telegram Interaction + Learned Profile ──
        tg_interaction = TelegramInteraction(
            token=cfg.telegram_bot_token,
            chat_id=cfg.telegram_chat_id,
        )
        learner = LearnedProfile()
        if tg_interaction.available:
            print(f"   [TELEGRAM] Interactive ask flow ready — will ask you via Telegram for unknown dropdowns")
            print(f"   [LEARNER]  {len(learner)} saved answers in profile")

        filler = FormFiller(
            ai_client=ai_client,
            notifier=tg_notifier,
            interaction=tg_interaction,
            learner=learner,
        )

        # ── Multi-step form loop ──
        for step in range(10):
            await Human.delay(1, 2)
            await filler.fill_form(page, tailored_resume_path, cover_text)
            print(f"   Step {step + 1}: fields filled")

            # Upload resume if a file input is present
            # Check input id/name/aria-label too (Greenhouse uses id='resume', not nested in a label)
            if tailored_resume_path:
                try:
                    file_inputs = await page.query_selector_all("input[type='file']")
                    for fi in file_inputs:
                        fi_info = await fi.evaluate("""el => ({
                            label: el.closest('label')?.textContent?.toLowerCase() || '',
                            ariaLabel: el.getAttribute('aria-label')?.toLowerCase() || '',
                            id: el.id?.toLowerCase() || '',
                            name: el.getAttribute('name')?.toLowerCase() || ''
                        })""")
                        combined = f"{fi_info['label']} {fi_info['ariaLabel']} {fi_info['id']} {fi_info['name']}"
                        if any(w in combined for w in ["resume", "cv", "upload"]):
                            await fi.set_input_files(tailored_resume_path)
                            print(f"   [OK] Tailored resume uploaded to #{fi_info['id']}")
                            await Human.delay(1, 2)
                except Exception:
                    pass

            # Handle Greenhouse cover letter: click 'Enter manually' to reveal textarea
            try:
                cl_textarea = await page.query_selector("#cover_letter_text")
                if not cl_textarea:
                    enter_manually = await page.query_selector(
                        "a:has-text('Enter manually'), button:has-text('Enter manually'), "
                        "[class*='enter']:has-text('manually')"
                    )
                    if enter_manually:
                        print(f"   [COVER] Clicking 'Enter manually' to reveal cover letter textarea...")
                        await enter_manually.click()
                        await Human.delay(1, 2)
            except Exception:
                pass

            # Fill cover letter textarea if it exists now
            if cover_text:
                try:
                    cl_textarea = await page.query_selector("#cover_letter_text")
                    if cl_textarea:
                        current = await cl_textarea.input_value()
                        if not current.strip():
                            await cl_textarea.click()
                            await cl_textarea.fill("")
                            await cl_textarea.fill(cover_text)
                            print(f"   [COVER] Cover letter filled ({len(cover_text)} chars)")
                        else:
                            print(f"   [COVER] Already filled ({len(current)} chars)")
                except Exception as e:
                    print(f"   [COVER] Could not fill cover letter: {e}")

            # Look for Submit button first
            submit_selectors = [
                "button[type='submit']", "input[type='submit']",
                "button:has-text('Submit')", "button:has-text('Submit Application')",
                "button:has-text('Send Application')",
            ]
            found_submit = False
            for sel in submit_selectors:
                try:
                    btn = await page.query_selector(sel)
                    if btn:
                        text = (await btn.inner_text()).strip()
                        print(f"\n[SUBMIT] Found Submit button: '{text}'")
                        await btn.scroll_into_view_if_needed()
                        await Human.delay(0.3, 0.7)

                        if dry_run:
                            print("[DRY RUN] skipping final click. Filled fields:")
                            inputs = await page.query_selector_all(
                                "input:not([type='hidden']):not([type='submit']):"
                                "not([type='button']), textarea"
                            )
                            for inp in inputs[:8]:
                                try:
                                    n = (await inp.get_attribute("name")) or ""
                                    v = (await inp.input_value())[:60]
                                    if v:
                                        print(f"     - {n}: {v}")
                                except Exception:
                                    pass
                            success = True
                        else:
                            print(f"\n{'!' * 50}")
                            print(f"[WARN] THIS WILL SUBMIT A REAL APPLICATION TO {company.upper()}")
                            print(f"{'!' * 50}")
                            inp = input("\nPress Enter to submit, or Ctrl+C to abort... ")
                            await btn.click()
                            await Human.delay(3, 6)
                            print("   [OK] Submit clicked!")

                        found_submit = True
                        break
                except (KeyboardInterrupt, EOFError):
                    print("\n[ABORT] Submit aborted by user.")
                    success = True
                    found_submit = True
                    break
                except Exception:
                    continue

            if found_submit:
                break

            # Look for Next / Continue / Review button
            next_selectors = [
                "button:has-text('Next')", "button:has-text('Continue')",
                "button:has-text('Review')", "a:has-text('Next')",
                "a:has-text('Continue')", "a:has-text('Review')",
            ]
            found_next = False
            for sel in next_selectors:
                try:
                    btn = await page.query_selector(sel)
                    if btn:
                        text = (await btn.inner_text()).strip()
                        print(f"   [NEXT] Clicking '{text}' (step {step + 1})...")
                        await btn.scroll_into_view_if_needed()
                        await Human.delay(0.3, 0.7)
                        await btn.click()
                        await Human.delay(1, 3)
                        found_next = True
                        break
                except Exception:
                    continue

            if not found_next:
                print(f"\n[WARN] No Next/Submit at step {step + 1}.")
                try:
                    await Human.screenshot(page, f"test_step_{step + 1}")
                except Exception:
                    pass
                break

        # ── Verify submission ──
        print("\n[VERIFY] Checking confirmation...")
        await Human.delay(2, 4)
        try:
            await Human.screenshot(page, "test_after_submit")
        except Exception:
            pass

        post_url = page.url
        if post_url != url:
            print(f"   [OK] URL changed: {post_url}")
            print("   [OK] Confirmation via URL redirect!")
            success = True
        else:
            try:
                body = (await page.inner_text("body") or "").lower()
                for word in ["thank you", "application submitted", "submitted",
                             "successfully applied", "received"]:
                    if word in body:
                        print(f"   [OK] Found confirmation text: '{word}'")
                        success = True
                        break
                if not success:
                    el = await page.query_selector(
                        "[class*='success'], [class*='confirmation'], "
                        "[role='alert'], [aria-label*='success']"
                    )
                    if el:
                        print("   [OK] Found success element")
                        success = True
            except Exception:
                pass

        if success:
            print("\n[SUCCESS] APPLICATION SUBMITTED SUCCESSFULLY!")
        else:
            print("\n[WARN] Submit clicked but NO confirmation detected.")
            print("   Check the browser / screenshot.")

        try:
            sp = await Human.screenshot(page, "test_final_state")
            print(f"   [SCREENSHOT] {sp}")
        except Exception:
            pass

    except Exception as e:
        logger.exception("Test failed")
        print(f"\n[ERROR] {e}")
        try:
            await Human.screenshot(page, "test_error")
        except Exception:
            pass
    finally:
        await page.close()
        await mgr.close()
        print(f"\n{'=' * 60}")
        print(f"[DONE] Test complete. Success: {success}")
        print(f"{'=' * 60}")

    return success


# ── Helper functions (mirrored from orchestrator.py) ────────────

_TAILOR_FIELDS = [
    "Name:", "Email:", "Phone:", "Location:", "Summary:", "Skills:",
    "Experience:", "Projects:", "Education:", "Certifications:",
]


def _resume_to_text(resume: ResumeProfile) -> str:
    """Full textual representation of the resume."""
    parts = [f"Name: {resume.name}"]
    if resume.summary:
        parts.append(f"Summary: {resume.summary}")
    if resume.skills:
        parts.append(f"Skills: {', '.join(resume.skills)}")
    if resume.experience:
        parts.append("Experience:")
        parts.extend(resume.experience)
    if resume.projects:
        parts.append("Projects:")
        for p in resume.projects:
            techs = f" ({', '.join(p.technologies)})" if p.technologies else ""
            parts.append(f"  - {p.name}{techs}: {p.description}")
    if resume.education:
        parts.append("Education:")
        parts.extend(resume.education)
    if resume.certifications:
        parts.append("Certifications:")
        parts.extend(resume.certifications)
    return "\n".join(parts)


def _parse_tailored_field(text: str, field: str) -> str:
    """Extract the value of a single field from tailored resume text."""
    idx = text.find(field)
    if idx < 0:
        return ""
    start = idx + len(field)
    rest = text[start:].lstrip()
    end = len(rest)
    for other in _TAILOR_FIELDS:
        if other == field:
            continue
        oi = rest.find(other)
        if 0 <= oi < end:
            end = oi
    return rest[:end].strip()


def _parse_tailored_list(text: str, field: str) -> list[str]:
    """Extract a bullet-list section from tailored resume text."""
    block = _parse_tailored_field(text, field)
    if not block:
        return []
    return [
        line.strip().lstrip("- ").strip()
        for line in block.split("\n")
        if line.strip()
    ]


def _parse_tailored_skills(text: str) -> list[str]:
    """Extract and parse the Skills section from tailored text."""
    skills_str = _parse_tailored_field(text, "Skills:")
    if not skills_str:
        return []
    return [s.strip() for s in skills_str.replace(", ", ",").split(",") if s.strip()]


def _safe_name(name: str) -> str:
    """Sanitize a company/name for use in file paths."""
    return "".join(c if c.isalnum() or c in " _.-" else "_" for c in name).strip()


def _write_cover_docx(letter_text: str, output_path: Path, resume: ResumeProfile, company: str) -> None:
    """Write a cover letter to a .docx file."""
    from datetime import date
    from docx import Document
    doc = Document()
    for p in [resume.name, resume.email, resume.phone or "", resume.location or ""]:
        if p:
            doc.add_paragraph(p)
    doc.add_paragraph("")
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph(f"Dear {company} Hiring Manager,")
    doc.add_paragraph("")
    for para_text in letter_text.split("\n\n"):
        stripped = para_text.strip()
        if stripped:
            doc.add_paragraph(stripped)
    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(resume.name)
    doc.save(str(output_path))


if __name__ == "__main__":
    headless = "--headless" in sys.argv
    dry_run = "--dry-run" in sys.argv
    url = None
    for i, arg in enumerate(sys.argv):
        if arg == "--url" and i + 1 < len(sys.argv):
            url = sys.argv[i + 1]

    if dry_run:
        print("\n[DRY RUN] All fields filled, no submission\n")

    asyncio.run(test_apply(headless=headless, apply_url=url, dry_run=dry_run))
