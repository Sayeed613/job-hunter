"""Resume document generator — creates styled DOCX and PDF files from a ResumeProfile."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from app.resume.models import ResumeProfile

logger = logging.getLogger("job_automation_bot")

# ── Style constants ───────────────────────────────────────────
_NAME_SIZE = Pt(20)
_SECTION_SIZE = Pt(14)
_BODY_SIZE = Pt(11)
_ACCENT_COLOR = RGBColor(0x1A, 0x56, 0xDB)  # Professional blue
_DARK_COLOR = RGBColor(0x33, 0x33, 0x33)


class ResumeGenerator:
    """Generates resume documents (DOCX and PDF) from a ResumeProfile.

    Usage:
        generator = ResumeGenerator()
        docx_path = generator.generate_docx(profile, "output/resume.docx")
        pdf_path = generator.generate_pdf(profile, "output/resume.pdf")
    """

    def __init__(self) -> None:
        self._has_libreoffice = self._check_libreoffice()

    # ── Public API ───────────────────────────────────────────

    def generate_docx(self, profile: ResumeProfile, output_path: str | Path) -> Path:
        """Generate a stylised DOCX resume file.

        Args:
            profile: The resume data to render.
            output_path: Filesystem path to write the .docx file.

        Returns:
            The Path to the generated file.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # -- Page margins --
        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # -- Name --
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(profile.name or "Candidate")
        name_run.bold = True
        name_run.font.size = _NAME_SIZE
        name_run.font.color.rgb = _DARK_COLOR

        # -- Contact info --
        contact_parts = []
        if profile.email:
            contact_parts.append(profile.email)
        if profile.phone:
            contact_parts.append(profile.phone)
        if profile.location:
            contact_parts.append(profile.location)
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run("  |  ".join(contact_parts))
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # -- Divider --
        self._add_divider(doc)

        # -- Summary --
        if profile.summary:
            self._add_section_header(doc, "PROFESSIONAL SUMMARY")
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(profile.summary)
            summary_run.font.size = _BODY_SIZE
            summary_run.font.color.rgb = _DARK_COLOR
            summary_para.paragraph_format.space_after = Pt(4)

        # -- Skills --
        if profile.skills:
            self._add_section_header(doc, "SKILLS")
            # Display skills as comma-separated with bullet-like formatting
            skills_text = ", ".join(profile.skills)
            skills_para = doc.add_paragraph()
            skills_run = skills_para.add_run(skills_text)
            skills_run.font.size = _BODY_SIZE
            skills_run.font.color.rgb = _DARK_COLOR
            skills_para.paragraph_format.space_after = Pt(4)

        # -- Experience --
        if profile.experience:
            self._add_section_header(doc, "EXPERIENCE")
            for exp in profile.experience:
                exp_para = doc.add_paragraph()
                exp_run = exp_para.add_run(exp)
                exp_run.font.size = _BODY_SIZE
                exp_run.font.color.rgb = _DARK_COLOR
                exp_para.paragraph_format.space_after = Pt(2)
                exp_para.paragraph_format.left_indent = Inches(0.25)

        # -- Projects --
        if profile.projects:
            self._add_section_header(doc, "PROJECTS")
            for proj in profile.projects:
                proj_para = doc.add_paragraph()
                proj_title_run = proj_para.add_run(proj.name)
                proj_title_run.bold = True
                proj_title_run.font.size = _BODY_SIZE
                proj_title_run.font.color.rgb = _ACCENT_COLOR
                if proj.technologies:
                    tech_run = proj_para.add_run(
                        f"  ({', '.join(proj.technologies)})"
                    )
                    tech_run.font.size = Pt(9)
                    tech_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                if proj.description:
                    proj_desc_para = doc.add_paragraph()
                    proj_desc_run = proj_desc_para.add_run(proj.description)
                    proj_desc_run.font.size = Pt(10)
                    proj_desc_run.font.color.rgb = _DARK_COLOR
                    proj_desc_para.paragraph_format.space_after = Pt(4)
                    proj_desc_para.paragraph_format.left_indent = Inches(0.25)

        # -- Education --
        if profile.education:
            self._add_section_header(doc, "EDUCATION")
            for edu in profile.education:
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(edu)
                edu_run.font.size = _BODY_SIZE
                edu_run.font.color.rgb = _DARK_COLOR
                edu_para.paragraph_format.space_after = Pt(2)
                edu_para.paragraph_format.left_indent = Inches(0.25)

        # -- Certifications --
        if profile.certifications:
            self._add_section_header(doc, "CERTIFICATIONS")
            for cert in profile.certifications:
                cert_para = doc.add_paragraph()
                cert_run = cert_para.add_run(cert)
                cert_run.font.size = _BODY_SIZE
                cert_run.font.color.rgb = _DARK_COLOR
                cert_para.paragraph_format.space_after = Pt(2)
                cert_para.paragraph_format.left_indent = Inches(0.25)

        doc.save(str(output_path))
        logger.info("Resume DOCX generated: %s", output_path)
        return output_path

    def generate_pdf(self, profile: ResumeProfile, output_path: str | Path) -> Optional[Path]:
        """Generate a PDF version of the resume using LibreOffice headless conversion.

        Falls back to returning None if LibreOffice is not available.

        Args:
            profile: The resume data to render.
            output_path: Filesystem path to write the .pdf file.

        Returns:
            The Path to the generated file, or None if conversion failed.
        """
        output_path = Path(output_path)
        if not self._has_libreoffice:
            logger.warning("LibreOffice not found — PDF generation disabled")
            return None

        try:
            import subprocess
            import tempfile

            # Generate DOCX first
            docx_path = output_path.with_suffix(".docx")
            self.generate_docx(profile, docx_path)

            # Convert to PDF using LibreOffice
            result = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_path.parent),
                    str(docx_path),
                ],
                capture_output=True,
                text=True,
                timeout=30,
            )

            if result.returncode != 0:
                logger.warning("LibreOffice PDF conversion failed: %s", result.stderr)
                return None

            if output_path.exists():
                logger.info("Resume PDF generated: %s", output_path)
                return output_path

            # LibreOffice may write to a different filename; try to find it
            generated = output_path.parent / f"{docx_path.stem}.pdf"
            if generated.exists():
                generated.rename(output_path)
                logger.info("Resume PDF generated: %s", output_path)
                return output_path

            return None
        except Exception as e:
            logger.warning("PDF generation failed: %s", e)
            return None

    # ── Helpers ───────────────────────────────────────────────

    @staticmethod
    def _add_section_header(doc: Document, title: str) -> None:
        """Add a styled section header with bottom border."""
        para = doc.add_paragraph()
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = _SECTION_SIZE
        run.font.color.rgb = _ACCENT_COLOR
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)

        # Add a thin bottom border via paragraph border
        pf = para.paragraph_format
        pPr = para._p.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        bottom = pBdr.find(qn("w:bottom"))
        if bottom is None:
            bottom = etree.SubElement(pBdr, qn("w:bottom"))
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1A56DB")

    @staticmethod
    def _add_divider(doc: Document) -> None:
        """Add a thin horizontal divider line."""
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pPr = para._p.get_or_add_pPr()
        pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        bottom = etree.SubElement(pBdr, qn("w:bottom"))
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")

    @staticmethod
    def _check_libreoffice() -> bool:
        """Check if LibreOffice is installed on the system."""
        try:
            import subprocess
            result = subprocess.run(
                ["soffice", "--headless", "--version"],
                capture_output=True,
                text=True,
                timeout=5,
            )
            return result.returncode == 0
        except Exception:
            return False
