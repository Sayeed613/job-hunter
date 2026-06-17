"""Async application pipeline — orchestrates the end-to-end job processing workflow."""

from __future__ import annotations

import asyncio
import logging
import random
from datetime import datetime, timedelta, timezone
from pathlib import Path

from app.ai.client import AIClient
from app.ai.cover_letter_gen import CoverLetterGenerator as AICoverLetterGen
from app.ai.keyword_extractor import KeywordExtractor
from app.ai.resume_tailor import ResumeTailor as AIResumeTailor
from app.browser.router import ApplicationRouter
from app.config.settings import Settings
from app.database.firestore_repository import FirestoreRepository
from app.models.application import Application
from app.models.job import Job
from app.resume.models import ResumeProfile
from app.tailor.resume_generator import ResumeGenerator
from app.notifier import LocalNotifier

logger = logging.getLogger("job_automation_bot")

# ── Tailored text field extractors ──────────────────────────
_TAILOR_FIELDS = [
    "Name:", "Email:", "Phone:", "Location:", "Summary:", "Skills:",
    "Experience:", "Projects:", "Education:", "Certifications:",
]

# ── Job title keywords to pre-filter relevant roles ─────────
# Expanded aggressively to catch every possible tech role
_RELEVANT_TITLE_KW = [
    "react", "python", "node", "frontend", "full stack", "fullstack",
    "backend", "back end", "typescript", "javascript", "developer",
    "engineer", "software", "web", "next", "api", "fastapi", "django",
    "vue", "angular", "sde", "swe", "programmer", "tech lead", "staff",
    "senior", "junior", "fresher", "associate", "full-stack",
    "front end", "back end", "mern", "mean", "ui", "ux",
    # Expanded keywords — catch more roles aggressively
    "tech", "coding", "program", "it ", "computer", "saas", "platform",
    "data", "analyst", "cloud", "devops", "sre", "infrastructure",
    "mobile", "app", "ios", "android", "flutter", "react native",
    "wordpress", "shopify", "cms", "automation", "support",
    "product", "qa", "test", "quality", "tester",
    "ai", "ml", "machine learning", "deep learning", "llm",
    "golang", "rust", "java", "spring", "c#", "dotnet",
    "graphql", "rest", "microservice", "distributed",
    "entry level", "entry-level", "0-1", "0 to 1",
]


def _parse_tailored_section(text: str, field: str) -> str:
    """Extract the value of a single field from tailored resume text."""
    idx = text.find(field)
    if idx < 0:
        return ""
    start = idx + len(field)
    rest = text[start:].lstrip()
    end = len(rest)
    for other in _TAILOR_FIELDS:
        if other == field:
            continue
        oi = rest.find(other)
        if 0 <= oi < end:
            end = oi
    return rest[:end].strip()


def _parse_tailored_list(text: str, field: str) -> list[str]:
    """Extract a bullet-list section from tailored resume text."""
    block = _parse_tailored_section(text, field)
    if not block:
        return []
    return [
        line.strip().lstrip("- ").strip()
        for line in block.split("\n")
        if line.strip()
    ]


def _parse_tailored_skills(text: str) -> list[str]:
    """Extract and parse the Skills section from tailored text."""
    skills_str = _parse_tailored_section(text, "Skills:")
    if not skills_str:
        return []
    return [s.strip() for s in skills_str.replace(", ", ",").split(",") if s.strip()]


def _job_title_matches(title: str) -> bool:
    """Check if a job title contains any relevant keyword (quick pre-filter)."""
    t = title.lower()
    return any(kw in t for kw in _RELEVANT_TITLE_KW)


class Pipeline:
    """Main application pipeline — search, tailor, apply, notify."""

    def __init__(
        self,
        ai_client: AIClient,
        repository: FirestoreRepository,
        notifier: LocalNotifier,
        settings: Settings,
    ) -> None:
        self._ai = ai_client
        self._repository = repository
        self._notifier = notifier
        self._settings = settings
        self._keyword_extractor = KeywordExtractor(client=ai_client)
        self._resume_tailor = AIResumeTailor(client=ai_client)
        self._cover_letter_gen = AICoverLetterGen(client=ai_client)
        self._resume_generator = ResumeGenerator()
        self._router = ApplicationRouter(
            ai_client=ai_client,
            notifier=notifier,
            interaction=None,
        )
        self._output_dir = Path("output")

    async def run_cycle(self, resume: ResumeProfile, providers: list) -> dict:
        await self._notifier.cycle_started(len(providers))

        # Reset auto-login attempt tracker so each cycle retries failed logins
        self._router.reset_login_attempts()

        # #region agent log
        import json as _json, time as _time
        with open("debug-eeb1f2.log", "a", encoding="utf-8") as _f:
            _f.write(_json.dumps({"sessionId": "eeb1f2", "hypothesisId": "A", "location": "orchestrator.py:run_cycle", "message": "ensure_browser_start", "data": {"headless": self._settings.headless, "provider_count": len(providers)}, "timestamp": int(_time.time() * 1000)}) + "\n")
        # #endregion
        try:
            await self._router.ensure_browser(
                headless=self._settings.headless,
                linkedin_email=self._settings.linkedin_email,
                linkedin_password=self._settings.linkedin_password,
            )
            # #region agent log
            with open("debug-eeb1f2.log", "a", encoding="utf-8") as _f:
                _f.write(_json.dumps({"sessionId": "eeb1f2", "hypothesisId": "A", "location": "orchestrator.py:run_cycle", "message": "ensure_browser_ok", "data": {"browser_ready": self._router.is_browser_ready}, "timestamp": int(_time.time() * 1000)}) + "\n")
            # #endregion
        except Exception as _browser_err:
            # #region agent log
            with open("debug-eeb1f2.log", "a", encoding="utf-8") as _f:
                _f.write(_json.dumps({"sessionId": "eeb1f2", "hypothesisId": "E", "location": "orchestrator.py:run_cycle", "message": "ensure_browser_failed", "data": {"error_type": type(_browser_err).__name__, "error": str(_browser_err)[:300]}, "timestamp": int(_time.time() * 1000)}) + "\n")
            # #endregion
            logger.warning(
                "Browser launch failed — continuing with API-only providers: %s",
                _browser_err,
            )

        if self._router.is_browser_ready:
            browser = self._router.get_browser()
            for p in providers:
                if hasattr(p, 'set_browser_manager'):
                    p.set_browser_manager(browser)
                    # Log which platforms have saved sessions
                    if p.requires_login and hasattr(browser, 'session'):
                        has = browser.session.has_session(p.platform)
                        if not has:
                            logger.info(
                                "%s needs login — run: python main.py --relogin %s",
                                p.name, p.platform,
                            )

        tasks = [asyncio.create_task(provider.fetch_jobs()) for provider in providers]
        # Timeout after 120 seconds — slow providers (LinkedIn, etc.) won't block the cycle
        done, pending = await asyncio.wait(tasks, timeout=120)
        for task in pending:
            task.cancel()

        all_jobs: list[Job] = []
        for task in done:
            try:
                r = task.result()
                if isinstance(r, list):
                    all_jobs.extend(r)
            except (asyncio.CancelledError, Exception):
                pass

        timed_out = len(pending)
        if timed_out:
            logger.info("%d provider(s) timed out after 120s — results so far: %d jobs", timed_out, len(all_jobs))

        logger.info("Collected %d total jobs from %d providers", len(all_jobs), len(providers))
        if not all_jobs:
            return {"found": 0, "applied": 0, "failed": 0, "skipped": 0}

        jobs = self._deduplicate(all_jobs)
        jobs, _filter_debug = self._filter_jobs(jobs)
        # #region agent log
        import json as _json2, time as _time2
        with open("debug-eeb1f2.log", "a", encoding="utf-8") as _f:
            _f.write(_json2.dumps({"sessionId": "eeb1f2", "hypothesisId": "C", "location": "orchestrator.py:run_cycle", "message": "filter_stats", "data": {"all_jobs": len(all_jobs), "after_dedup": len(jobs) + _filter_debug.get("filtered_out", 0), "after_filter": len(jobs), **_filter_debug}, "timestamp": int(_time2.time() * 1000)}) + "\n")
        # #endregion

        new_jobs: list[Job] = []
        for job in jobs:
            existing = self._repository.get_application(job.job_id)
            if not existing:
                new_jobs.append(job)

        # ── Pre-filter: only keep jobs with relevant titles ──────────
        relevant_jobs = [j for j in new_jobs if _job_title_matches(j.title)]
        skipped_due_to_title = len(new_jobs) - len(relevant_jobs)
        if skipped_due_to_title:
            logger.info("Title pre-filter: removed %d non-tech jobs, keeping %d", skipped_due_to_title, len(relevant_jobs))

        # ── Sort by relevance: score ALL jobs and take top N ─────────
        full_resume_text = self._resume_to_text(resume)
        scored_jobs: list[tuple[float, Job]] = []
        for job in relevant_jobs:
            score = self._quick_score(full_resume_text, job.description)
            # Title match bonus
            title_lower = job.title.lower()
            for kw in resume.skills:
                if kw.lower() in title_lower:
                    score += 0.05
            # Short description penalty — less data to match = less reliable
            word_count = len(job.description.split()) if job.description else 0
            if word_count < 20:
                score *= 0.5
            scored_jobs.append((score, job))

        scored_jobs.sort(key=lambda x: x[0], reverse=True)
        new_jobs = [job for _, job in scored_jobs]

        await self._notifier.jobs_found(len(new_jobs), len(all_jobs))
        if not new_jobs:
            logger.info(
                "No relevant jobs after title filter (%d total, %d passed location filter)",
                len(all_jobs), len(jobs),
            )
            return {"found": len(all_jobs), "applied": 0, "failed": 0, "skipped": len(all_jobs)}

        applied = 0
        failed = 0
        skipped = 0
        max_apply = min(self._settings.max_applications_per_cycle, len(new_jobs))

        for i, job in enumerate(new_jobs[:max_apply]):
            try:
                await self._notifier.job_processing(
                    i + 1, max_apply, job.title, job.company,
                    job.location, job.remote_type, job.salary or "",
                    job.apply_url,
                )

                # ── AI keyword extraction ────────────────────────
                ai_matched = 0
                try:
                    keywords = await self._keyword_extractor.extract(job.description)
                    jd_keyword_count = len(keywords.get("hard_skills", []))
                    ai_matched = sum(
                        1 for s in keywords.get("hard_skills", [])
                        if s.lower() in " ".join(resume.skills).lower()
                    )
                    await self._notifier.tailoring(ai_matched, jd_keyword_count)
                except Exception:
                    logger.warning("AI keyword extraction failed — using quick score only")
                    keywords = {"hard_skills": [], "soft_skills": [], "years_required": 0}
                    jd_keyword_count = 0
                    ai_matched = 0

                # ── Match score check (AGGRESSIVE) ───────────────
                match_score = self._quick_score(full_resume_text, job.description)

                # AGGRESSIVE MODE: accept almost anything that passed the title filter
                # Accept if AI matched 1+ hard skills, OR quick_score is above a very low threshold
                min_ai = self._settings.min_ai_skills
                min_score = self._settings.min_match_score
                if ai_matched < min_ai and match_score < min_score:
                    logger.info(
                        "Skipping %s — score %.2f (AI %d/%d) below thresholds (%d / %.2f)",
                        job.title, match_score, ai_matched, jd_keyword_count,
                        min_ai, min_score,
                    )
                    skipped += 1
                    continue

                logger.info(
                    "Matched %s — score: %.2f, AI skills: %d/%d",
                    job.title, match_score, ai_matched, jd_keyword_count,
                )

                # ── Resume tailoring ─────────────────────────────
                base_resume_text = self._resume_to_text(resume)
                if self._ai.is_available:
                    try:
                        tailored_text = await self._resume_tailor.tailor(
                            base_resume_text, job, keywords.get("hard_skills", []),
                        )
                    except Exception:
                        logger.warning("AI resume tailoring failed — using base resume")
                        tailored_text = base_resume_text
                else:
                    tailored_text = base_resume_text

                tailored_profile = ResumeProfile(
                    name=_parse_tailored_section(tailored_text, "Name:") or resume.name,
                    email=_parse_tailored_section(tailored_text, "Email:") or resume.email,
                    phone=_parse_tailored_section(tailored_text, "Phone:") or resume.phone,
                    location=_parse_tailored_section(tailored_text, "Location:") or resume.location,
                    summary=_parse_tailored_section(tailored_text, "Summary:") or resume.summary,
                    skills=_parse_tailored_skills(tailored_text) or resume.skills,
                    experience=_parse_tailored_list(tailored_text, "Experience:") or resume.experience,
                    projects=resume.projects,
                    education=_parse_tailored_list(tailored_text, "Education:") or resume.education,
                    certifications=_parse_tailored_list(tailored_text, "Certifications:") or resume.certifications,
                )

                # ── Generate resume DOCX ────────────────────────
                job_dir = self._output_dir / _safe_name(job.company)
                job_dir.mkdir(parents=True, exist_ok=True)
                safe_company = _safe_name(job.company)
                resume_docx = self._resume_generator.generate_docx(
                    tailored_profile, job_dir / f"sayeed_ahmed-{safe_company}.docx",
                )
                try:
                    self._resume_generator.generate_pdf(
                        tailored_profile, job_dir / f"sayeed_ahmed-{safe_company}.pdf",
                    )
                except RuntimeError:
                    pass
                resume_path = str(resume_docx)

                # ── Cover letter ────────────────────────────────
                ai_cover_letter = False
                if self._ai.is_available:
                    try:
                        cl_summary = (
                            _parse_tailored_section(tailored_text, "Summary:")
                            or resume.summary
                        )
                        cover_letter_text = await self._cover_letter_gen.generate(
                            cl_summary,
                            [f"- {p.description}" for p in resume.projects[:5]],
                            job, keywords.get("hard_skills", []),
                        )
                        ai_cover_letter = True
                    except Exception:
                        logger.warning("AI cover letter failed — using template")
                        cover_letter_text = (
                            f"Dear {job.company} Hiring Team,\n\n"
                            f"I am excited to apply for the {job.title} position. "
                            f"With my background in frontend development and a passion "
                            f"for building great user interfaces, I believe I would be "
                            f"a strong addition to your team.\n\n"
                            f"Thank you for your consideration.\n\n"
                            f"Best regards,\n{resume.name}"
                        )
                else:
                    cover_letter_text = (
                        f"Dear {job.company} Hiring Team,\n\n"
                        f"I am excited to apply for the {job.title} position.\n\n"
                        f"Best regards,\n{resume.name}"
                    )
                cover_path = job_dir / f"cover_letter_{job.job_id}.docx"
                self._write_docx(cover_letter_text, cover_path, resume, job.company)
                cover_letter_path = str(cover_path)

                # ── Submit application ──────────────────────────
                await self._notifier.applying("browser")
                success = await self._router.apply(
                    job, resume_path, cover_letter_text, cover_letter_path,
                )

                app = Application(
                    job_id=job.job_id, title=job.title, company=job.company,
                    location=job.location, remote_type=job.remote_type,
                    job_type=job.job_type, salary=job.salary, source=job.source,
                    apply_url=job.apply_url, posted_at=job.posted_at,
                    applied_at=datetime.now(timezone.utc),
                    status="applied" if success else "failed",
                    application_method="browser", resume_path=resume_path,
                    cover_letter_path=cover_letter_path,
                    matched_keywords=keywords.get("hard_skills", []),
                    match_score=match_score,
                )
                self._repository.save_application(app)

                if success:
                    await self._notifier.success(
                        job.title, job.company,
                        resume_path=resume_path,
                        cover_letter_path=cover_letter_path if ai_cover_letter else "",
                        salary=job.salary or "",
                        location=job.location,
                    )
                    applied += 1
                else:
                    await self._notifier.failure(
                        job.title, job.company, "Form submission failed", job.apply_url,
                    )
                    failed += 1

                # AGGRESSIVE MODE: shorter delays between applications
                await asyncio.sleep(random.uniform(12, 25))
            except Exception as e:
                logger.exception("Failed to process job %s", job.job_id)
                await self._notifier.failure(job.title, job.company, str(e), job.apply_url)
                failed += 1

        await self._router.close_browser()
        next_run = f"{self._settings.run_interval_hours} hours"
        await self._notifier.cycle_summary(applied, failed, skipped, next_run)
        logger.info(
            "Cycle complete: %d applied, %d failed, %d skipped",
            applied, failed, skipped,
        )
        _result = {
            "found": len(all_jobs),
            "applied": applied,
            "failed": failed,
            "skipped": skipped + (len(new_jobs) - max_apply),
        }
        # #region agent log
        import json as _json3, time as _time3
        with open("debug-eeb1f2.log", "a", encoding="utf-8") as _f:
            _f.write(_json3.dumps({"sessionId": "eeb1f2", "hypothesisId": "D", "location": "orchestrator.py:run_cycle", "message": "cycle_complete", "data": _result, "timestamp": int(_time3.time() * 1000)}) + "\n")
        # #endregion
        return _result

    @staticmethod
    def _deduplicate(jobs: list[Job]) -> list[Job]:
        seen_urls: set[str] = set()
        seen_ct: set[str] = set()
        unique: list[Job] = []
        for job in jobs:
            url_key = job.apply_url.lower().strip()
            if url_key and url_key in seen_urls:
                continue
            seen_urls.add(url_key)
            ct_key = f"{job.company.lower().strip()}:{job.title.lower().strip()}"
            if ct_key in seen_ct:
                continue
            seen_ct.add(ct_key)
            unique.append(job)
        return unique

    def _filter_jobs(self, jobs: list[Job]) -> tuple[list[Job], dict]:
        """Filter jobs: remote anywhere in the world + Bangalore onsite/hybrid.

        Rules:
        1. Remote jobs → apply anywhere in the world.
        2. Bangalore jobs → apply even if onsite/hybrid (user can commute).
        3. Everything else → skip (non-remote, non-Bangalore).
        4. Time: only jobs posted within the last N hours (default 48).
        5. Experience: within min/max range.
        """
        excluded = self._get_excluded_companies()
        now = datetime.now(timezone.utc)
        max_age = timedelta(hours=getattr(self._settings, "max_job_age_hours", 48))
        bangalore_keywords = ("bangalore", "bengaluru")
        filtered: list[Job] = []
        hybrid_non_bangalore_leaked = 0
        filtered_out = 0

        for job in jobs:
            if job.company.lower().strip() in excluded:
                continue

            rt = job.remote_type.lower()
            loc = job.location.lower()

            is_remote = (
                "remote" in rt
                or rt == ""
                or "remote" in loc
                or "anywhere" in loc
                or "global" in loc
                or "work from home" in loc
                or "wfh" in loc
            )
            is_bangalore = any(k in loc for k in bangalore_keywords)

            if not is_remote and not is_bangalore:
                filtered_out += 1
                continue

            if job.posted_at is not None:
                posted = job.posted_at
                if posted.tzinfo is None:
                    posted = posted.replace(tzinfo=timezone.utc)
                age = now - posted
                if age > max_age:
                    continue

            if job.experience_years is not None:
                if job.experience_years < self._settings.min_experience:
                    continue
                if job.experience_years > self._settings.max_experience:
                    continue

            filtered.append(job)

        logger.info(
            "Filter: %d jobs in, %d out (remote worldwide + Bangalore)",
            len(jobs), len(filtered),
        )
        return filtered, {
            "filtered_out": filtered_out,
            "hybrid_non_bangalore_leaked": hybrid_non_bangalore_leaked,
        }

    def _get_excluded_companies(self) -> set[str]:
        raw = self._settings.excluded_companies
        if not raw:
            return set()
        return {c.strip().lower() for c in raw.split(",") if c.strip()}

    @staticmethod
    def _quick_score(resume_text: str, job_description: str) -> float:
        """Word-overlap score between full resume text and job description."""
        resume_lower = resume_text.lower()
        jd_lower = job_description.lower()
        resume_words = set(w for w in resume_lower.split() if len(w) > 3)
        jd_words = set(w for w in jd_lower.split() if len(w) > 3)
        if not jd_words:
            return 0.5
        overlap = resume_words & jd_words
        return len(overlap) / len(jd_words)

    @staticmethod
    def _resume_to_text(resume: ResumeProfile) -> str:
        """Full textual representation of the resume for matching."""
        parts = [f"Name: {resume.name}"]
        if resume.summary:
            parts.append(f"Summary: {resume.summary}")
        if resume.skills:
            parts.append(f"Skills: {', '.join(resume.skills)}")
        if resume.experience:
            parts.append("Experience:")
            parts.extend(resume.experience)
        if resume.projects:
            parts.append("Projects:")
            for p in resume.projects:
                techs = f" ({', '.join(p.technologies)})" if p.technologies else ""
                parts.append(f"  - {p.name}{techs}: {p.description}")
        if resume.education:
            parts.append("Education:")
            parts.extend(resume.education)
        if resume.certifications:
            parts.append("Certifications:")
            parts.extend(resume.certifications)
        return "\n".join(parts)

    @staticmethod
    def _write_docx(letter_text: str, output_path: Path, resume: ResumeProfile, company: str) -> None:
        from datetime import date
        from docx import Document
        doc = Document()
        for p in [resume.name, resume.email, resume.phone or "", resume.location or ""]:
            if p:
                doc.add_paragraph(p)
        doc.add_paragraph("")
        doc.add_paragraph(date.today().strftime("%B %d, %Y"))
        doc.add_paragraph(f"Dear {company} Hiring Manager,")
        doc.add_paragraph("")
        for para_text in letter_text.split("\n\n"):
            stripped = para_text.strip()
            if stripped:
                doc.add_paragraph(stripped)
        doc.add_paragraph("")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph(resume.name)
        doc.save(str(output_path))


def _safe_name(name: str) -> str:
    return "".join(c if c.isalnum() or c in " _.-" else "_" for c in name).strip()
