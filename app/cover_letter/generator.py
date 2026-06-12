"""Cover letter generator — creates personalised cover letters using AI."""

from __future__ import annotations

import logging
import os
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.ai.opencode_client import OpenCodeClient
from app.models.job import Job
from app.resume.models import ResumeProfile

logger = logging.getLogger("headhunter")

# ── Cover-letter prompt ──────────────────────────────────────

_COVER_LETTER_PROMPT = """You are a professional job application writer. Write a compelling, personalised cover letter for a job application.

Use the candidate's resume and the job description to write a letter that:

1. Opens with a strong introduction expressing interest in the specific role and company.
2. Highlights the candidate's most relevant skills and experience that match the job requirements.
3. Mentions specific projects that demonstrate relevant expertise: {projects}.
4. Explains why the candidate is a good fit for this particular role and company.
5. Closes professionally with enthusiasm and a call to action.

## Important rules

- Be professional and confident but not arrogant.
- Be specific — reference technologies, tools, and experiences from the resume.
- Do NOT invent experience, skills, or qualifications that are not present in the resume.
- Keep it between 250 and 400 words.
- Address the hiring manager professionally (use "Hiring Manager" if no name is known).
- Output ONLY the letter body — no subject line, no salutation placeholder, no instructions.

## Job Posting

**Title:** {title}
**Company:** {company}
**Description:**
{description}

## Candidate Resume

**Name:** {name}
**Summary:** {summary}
**Skills:** {skills}
**Experience:** {experience}
**Projects:** {projects_detail}
**Education:** {education}"""


class CoverLetterGenerator:
    """Generates personalised cover letters using an LLM.

    Produces both plain-text and formatted DOCX output.
    """

    def __init__(
        self,
        client: OpenCodeClient,
        temperature: float = 0.4,
        max_tokens: int = 600,
    ) -> None:
        """Initialise the generator.

        Args:
            client: An :class:`OpenCodeClient` instance.
            temperature: Sampling temperature (default 0.4).
            max_tokens: Maximum tokens in the response (default 600).
        """
        self._client = client
        self._temperature = temperature
        self._max_tokens = max_tokens

    # ── Public API ───────────────────────────────────────────

    def generate(
        self,
        job: Job,
        resume: ResumeProfile,
        selected_projects: list[str] | None = None,
    ) -> str:
        """Generate a plain-text cover letter.

        Args:
            job: The :class:`Job` to apply for.
            resume: The candidate's :class:`ResumeProfile`.
            selected_projects: Optional list of project names to
                highlight.  When omitted, all projects from the resume
                are included.

        Returns:
            The cover letter body as a plain string.
        """
        prompt = self._build_prompt(job, resume, selected_projects)

        try:
            letter = self._client.chat(
                prompt,
                temperature=self._temperature,
                max_tokens=self._max_tokens,
            )
        except Exception:
            logger.exception(
                "Failed to generate cover letter for %s at %s",
                job.title, job.company,
            )
            raise RuntimeError(
                f"Cover letter generation failed for {job.title} at {job.company}"
            ) from None

        letter = letter.strip().strip('"').strip("'")

        logger.info(
            "Cover letter generated",
            extra={
                "company": job.company,
                "role": job.title,
                "length": len(letter),
            },
        )

        return letter

    def generate_to_docx(
        self,
        job: Job,
        resume: ResumeProfile,
        output_path: str | Path,
        selected_projects: list[str] | None = None,
        company_name: str | None = None,
        hiring_manager: str | None = None,
    ) -> Path:
        """Generate a cover letter and save it as a .docx file.

        Args:
            job: The :class:`Job` to apply for.
            resume: The candidate's :class:`ResumeProfile`.
            output_path: Destination file path.
            selected_projects: Optional list of project names to
                highlight.
            company_name: Optional company name override for the
                letter header.
            hiring_manager: Optional hiring manager name for the
                salutation.

        Returns:
            The resolved :class:`Path` of the written file.
        """
        letter_text = self.generate(job, resume, selected_projects)
        return self._write_docx(
            letter_text, output_path, resume, company_name or job.company,
            hiring_manager,
        )

    # ── Prompt builder ───────────────────────────────────────

    def _build_prompt(
        self,
        job: Job,
        resume: ResumeProfile,
        selected_projects: list[str] | None = None,
    ) -> str:
        """Build the LLM prompt with job and resume context."""
        projects_list = selected_projects or [
            p.name for p in resume.projects
        ]
        projects_str = ", ".join(projects_list) if projects_list else "None listed"

        projects_detail_lines = []
        for p in resume.projects:
            tech_str = f" ({', '.join(p.technologies)})" if p.technologies else ""
            projects_detail_lines.append(f"- {p.name}{tech_str}: {p.description or 'No description'}")
        projects_detail = "\n".join(projects_detail_lines) if projects_detail_lines else "No projects listed"

        return _COVER_LETTER_PROMPT.format(
            title=job.title,
            company=job.company,
            description=job.description,
            name=resume.name,
            summary=resume.summary,
            skills=", ".join(resume.skills),
            experience="\n".join(f"- {exp}" for exp in resume.experience),
            projects=projects_str,
            projects_detail=projects_detail,
            education="\n".join(f"- {edu}" for edu in resume.education),
        )

    # ── DOCX writer ──────────────────────────────────────────

    @staticmethod
    def _write_docx(
        letter_text: str,
        output_path: str | Path,
        resume: ResumeProfile,
        company_name: str,
        hiring_manager: str | None = None,
    ) -> Path:
        """Write the cover letter text into a formatted .docx file."""
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # ── Sender info ──────────────────────────────────────
        sender_info = [resume.name, resume.email]
        if resume.phone:
            sender_info.append(resume.phone)
        if resume.location:
            sender_info.append(resume.location)

        for line in sender_info:
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            for run in para.runs:
                run.font.size = Pt(11)

        # ── Date ─────────────────────────────────────────────
        date_para = doc.add_paragraph(date.today().strftime("%B %d, %Y"))
        date_para.paragraph_format.space_before = Pt(12)
        date_para.paragraph_format.space_after = Pt(12)

        # ── Recipient ────────────────────────────────────────
        if hiring_manager:
            doc.add_paragraph(f"Dear {hiring_manager},")
        elif company_name:
            doc.add_paragraph(f"Dear {company_name} Hiring Manager,")
        else:
            doc.add_paragraph("Dear Hiring Manager,")

        # ── Body ─────────────────────────────────────────────
        for para_text in letter_text.split("\n\n"):
            stripped = para_text.strip()
            if not stripped:
                continue
            body_para = doc.add_paragraph(stripped)
            body_para.paragraph_format.space_before = Pt(6)
            body_para.paragraph_format.space_after = Pt(6)
            body_para.paragraph_format.line_spacing = Pt(14)
            for run in body_para.runs:
                run.font.size = Pt(11)

        # ── Closing ──────────────────────────────────────────
        doc.add_paragraph("")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph(resume.name)

        doc.save(str(output))

        logger.info(
            "Cover letter DOCX saved",
            extra={
                "path": str(output),
                "size_bytes": os.path.getsize(output),
            },
        )

        return output
